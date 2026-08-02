from pathlib import Path
from shutil import copy2

from docx import Document


BASE = Path(r"C:\Users\Roberto\OneDrive - artelnics.com\DIPCAN\subsanacion_tecnica")
DOCX = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO.docx"
BACKUP = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO_pre_recuento_radiomica.bak.docx"


def paragraph_starting(doc, prefix):
    matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]


def main():
    copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    paragraph_starting(
        doc, "De las características radiómicas originales"
    ).text = (
        "Quibim realizó una extracción bruta de 1.018 características radiómicas por lesión. "
        "Posteriormente, durante el curado y la agregación de la información por paciente, este "
        "conjunto se transformó en 127 variables radiómicas curadas, cifra registrada en "
        "dashboard.xlsx. Sobre estas variables se aplicó el proceso de selección para reducir "
        "redundancias y correlaciones antes del entrenamiento del modelo."
    )

    paragraph_starting(
        doc, "La siguiente tabla muestra el número de características radiómicas originales"
    ).text = (
        "La siguiente tabla diferencia las tres etapas del tratamiento de la información "
        "radiómica: extracción bruta por lesión, variables curadas y variables finales "
        "utilizadas como entradas del modelo."
    )

    paragraph_starting(
        doc, "Tras el curado y la agregación de la información por paciente"
    ).text = (
        "La extracción inicial realizada por Quibim produjo 1.018 características radiómicas "
        "brutas por lesión. Tras el curado y la agregación de la información por paciente, "
        "dashboard.xlsx registra 127 variables radiómicas curadas. El análisis de redundancia "
        "y correlación redujo posteriormente este conjunto a 74 variables empleadas como "
        "entradas del modelo, tal como se recoge en la sección de métodos."
    )

    target = None
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Características radiómicas originales":
            target = table
            break
    if target is None:
        raise RuntimeError("Radiomics count table not found")

    target.rows[0].cells[0].text = "Características radiómicas brutas extraídas por lesión"
    target.rows[0].cells[1].text = "1.018"
    target.rows[1].cells[0].text = "Variables radiómicas curadas (dashboard.xlsx)"
    target.rows[1].cells[1].text = "127"
    row = target.add_row()
    row.cells[0].text = "Variables radiómicas finales utilizadas en el modelo"
    row.cells[1].text = "74"

    doc.core_properties.comments = (
        "Se distinguen tres etapas: 1.018 características brutas por lesión según Quibim, "
        "127 variables radiómicas curadas según dashboard.xlsx y 74 entradas finales del modelo."
    )
    doc.save(DOCX)
    print(DOCX)
    print(BACKUP)


if __name__ == "__main__":
    main()
