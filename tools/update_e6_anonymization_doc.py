from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"C:\Users\Roberto\OneDrive - artelnics.com\DIPCAN\subsanacion_tecnica")
SOURCE = BASE / "Subsanacion_E6_Script_Anonimizacion_Artelnics_TSI-100206-2021-2.docx"
BACKUP = BASE / "Subsanacion_E6_Script_Anonimizacion_Artelnics_TSI-100206-2021-2_pre_script_original.bak.docx"
OUTPUT = BASE / "Subsanacion_E6_Script_Anonimizacion_Artelnics_TSI-100206-2021-2_actualizado.docx"

BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"
RED = "9C0006"
GREEN = "006100"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        c = hdr.cells[i]
        c.text = text
        shade(c, BLUE)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cells[i])
            if ri % 2:
                shade(cells[i], LIGHT_GREY)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(8.5)
        if widths:
            for c, width in zip(cells, widths):
                c.width = Cm(width)
    doc.add_paragraph()
    return table


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F7F7F7")
    margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for line_no, line in enumerate(text.rstrip().splitlines(), 1):
        r = p.add_run(f"{line_no:>2}  {line}")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8)
        if line_no != len(text.rstrip().splitlines()):
            r.add_break()
    doc.add_paragraph()


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_note(doc, title, text, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, color)
    margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + " ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph()


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.04
    for name, size, color in (
        ("Title", 20, BLUE),
        ("Subtitle", 12, "4F4F4F"),
        ("Heading 1", 15, BLUE),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10.5, BLUE),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.text = "PROYECTO DIPCAN  |  SUBSANACIÓN TÉCNICA E6"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("666666")
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Artificial Intelligence Techniques, S.L.  ·  TSI-100206-2021-2  ·  ")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def main():
    if not BACKUP.exists():
        copy2(SOURCE, BACKUP)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    add_header_footer(section)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PROYECTO DIPCAN")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Respuesta al requerimiento de la visita técnica")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Entregable E6 — Script para anonimizar datos personales de los informes radiómicos durante el curado de la base de datos")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    add_table(
        doc,
        ["Dato", "Referencia"],
        [
            ("Convocatoria", "Misiones de I+D en Inteligencia Artificial 2021"),
            ("Proyecto", "DIPCAN — Digitalización y Manejo Integral de la Medicina Personalizada"),
            ("Expediente", "TSI-100206-2021-2"),
            ("Entidad beneficiaria", "Artificial Intelligence Techniques, S.L. (Artelnics) — NIF B37536588"),
            ("Entregable", "E6 — Base de datos curada para generación del algoritmo"),
            ("Visita técnica", "22 de julio de 2026, telemática"),
            ("Documento de referencia", "Justificación de visita técnica, expedientes TSI-100206-2021-2, -3 y -7"),
            ("Fecha de actualización", "28 de julio de 2026"),
        ],
    )

    doc.add_heading("1. Objeto y respuesta aportada", level=1)
    doc.add_paragraph(
        "El documento de justificación de la visita técnica requiere literalmente: "
        "«Script para anonimizar los datos personales de los informes radiómicos durante el "
        "proceso de curado de base de datos, correspondiente al entregable 6»."
    )
    doc.add_paragraph(
        "En respuesta, se aporta el archivo original dicom_anonymizer.py recuperado de la "
        "carpeta de trabajo del proyecto, junto con sus archivos de ejemplo example.dcm y "
        "anonymized_example.DCM. Este documento reproduce el código sin ampliarlo ni atribuirle "
        "funciones que no contiene, explica su operación y presenta comprobaciones realizadas "
        "exclusivamente sobre el ejemplo incluido con el propio script."
    )
    add_note(
        doc,
        "Precisión terminológica.",
        "El script actúa sobre los metadatos de los archivos DICOM que sirven de entrada al "
        "análisis radiómico. Las variables radiómicas derivadas son magnitudes cuantitativas; "
        "por ello, la retirada del nombre se realiza antes de su extracción.",
    )

    doc.add_heading("1.1. Desarrollo y validación del procedimiento", level=2)
    doc.add_paragraph(
        "Durante las actividades técnicas del proyecto, Artelnics desarrolló un script en Python "
        "para anonimizar archivos de imagen médica en formato DICOM, utilizando la biblioteca "
        "pydicom[1]. El procedimiento recibía un archivo DICOM de entrada, modificaba o eliminaba "
        "los campos identificativos definidos y generaba un nuevo archivo de salida anonimizado. "
        "Un ejemplo simplificado del proceso implementado se reproduce íntegramente en el apartado siguiente."
    )
    doc.add_paragraph(
        "El funcionamiento del script se comprobó mediante imágenes de prueba proporcionadas durante "
        "el proyecto. Estas pruebas permitieron verificar la correcta lectura, modificación y escritura "
        "de los metadatos DICOM, así como evaluar si era necesario incorporar campos adicionales al "
        "proceso de anonimización."
    )

    doc.add_heading("2. Código fuente original", level=1)
    doc.add_paragraph(
        "Contenido íntegro del archivo dicom_anonymizer.py (fecha interna: 7 de marzo de 2022; autor: Artelnics):"
    )
    doc.add_paragraph(
        "El programa importa pydicom, abre el archivo example.DCM y carga su conjunto de datos. "
        "A continuación elimina todas las etiquetas privadas mediante remove_private_tags(), "
        "sustituye el valor de PatientName (0010,0010) por la cadena ANONYMIZED y guarda el resultado "
        "como anonymized_example.DCM. La información de píxel se conserva porque el código no la "
        "decodifica, transforma ni reasigna."
    )
    code = '''# -*- coding: utf-8 -*-
"""
Created on Mon Mar  7 12:44:00 2022

@author: Artelnics
"""

import pydicom

fpath = "example.DCM"

dicom = pydicom.dcmread(fpath)

dicom.remove_private_tags()

dicom.data_element("PatientName").value = "ANONYMIZED";

dicom.save_as('anonymized_example.DCM')'''
    add_code(doc, code)

    doc.add_heading("3. Funcionamiento y modo de ejecución", level=1)
    add_table(
        doc,
        ["Paso", "Operación real del script", "Resultado"],
        [
            ("1", 'Define como entrada el archivo fijo "example.DCM".', "Selecciona un único DICOM."),
            ("2", "Lee el archivo mediante pydicom.dcmread().", "Carga metadatos y píxeles."),
            ("3", "Ejecuta remove_private_tags().", "Elimina todos los elementos de grupos privados."),
            ("4", 'Asigna "ANONYMIZED" a PatientName (0010,0010).', "Sustituye el nombre del paciente."),
            ("5", 'Guarda "anonymized_example.DCM".', "Genera una copia procesada."),
        ],
    )
    doc.add_paragraph("Requisitos y ejecución desde la carpeta del script:")
    add_code(doc, "python -m pip install pydicom\npython dicom_anonymizer.py")
    add_bullet(doc, "La ruta de entrada y la de salida están codificadas en el propio archivo.")
    add_bullet(doc, "El script procesa un archivo por ejecución y no incluye interfaz de línea de comandos.")
    add_bullet(doc, "En Windows, la diferencia entre example.DCM y example.dcm no impide la lectura; en sistemas sensibles a mayúsculas debe igualarse el nombre.")

    doc.add_heading("4. Ejemplo incluido con el script: antes y después", level=1)
    doc.add_paragraph(
        "Se compararon los archivos example.dcm y anonymized_example.DCM sin exponer sus valores. "
        "La tabla refleja únicamente presencia, longitud o estado del campo."
    )
    add_table(
        doc,
        ["Elemento DICOM", "Entrada", "Salida", "Efecto comprobado"],
        [
            ("PatientName (0010,0010)", "Presente (21 caracteres)", "ANONYMIZED (10 caracteres)", "Sustituido"),
            ("Elementos privados", "179", "0", "Eliminados"),
            ("PatientID (0010,0020)", "Presente (4 caracteres)", "Presente (4 caracteres)", "Sin cambio"),
            ("PatientBirthDate (0010,0030)", "Vacío", "Vacío", "Sin cambio"),
            ("PatientSex (0010,0040)", "Presente", "Presente", "Sin cambio"),
            ("StudyDate (0008,0020)", "Presente", "Presente", "Sin cambio"),
            ("InstitutionName (0008,0080)", "Presente", "Presente", "Sin cambio"),
            ("Study/Series/SOP Instance UID", "Presentes", "Presentes", "Sin cambio"),
            ("PixelData (7FE0,0010)", "32.768 bytes", "32.768 bytes", "Idéntico byte a byte"),
        ],
    )
    add_note(
        doc,
        "Resultado del ejemplo.",
        "El nombre queda sustituido y se eliminan las etiquetas privadas, mientras que los datos "
        "de imagen permanecen intactos. Esto permite continuar el análisis radiológico/radiómico "
        "sin recodificar los píxeles.",
        color="E2F0D9",
    )
    doc.add_heading("4.1. Representación JSON de la comprobación", level=2)
    doc.add_paragraph(
        "El siguiente JSON resume algunos campos clave del par de prueba. El valor original de "
        "PatientName se omite deliberadamente; «sin cambios» significa que el script conserva el "
        "valor existente, no que lo sustituya por ese texto."
    )
    json_example = '''{
  "entrada": {
    "archivo": "example.dcm",
    "PatientName": "[valor original omitido]",
    "PatientID": "[sin cambios]",
    "PatientBirthDate": "",
    "private_tags": 179,
    "PixelData_bytes": 32768
  },
  "salida": {
    "archivo": "anonymized_example.DCM",
    "PatientName": "ANONYMIZED",
    "PatientID": "[sin cambios]",
    "PatientBirthDate": "",
    "private_tags": 0,
    "PixelData_bytes": 32768
  },
  "verificacion": {
    "PatientName_sustituido": true,
    "etiquetas_privadas_eliminadas": true,
    "PixelData_identico_byte_a_byte": true
  }
}'''
    add_code(doc, json_example)

    doc.add_heading("5. Alcance probado y limitaciones", level=1)
    doc.add_paragraph(
        "La evidencia disponible permite afirmar que el script original retira el nombre del "
        "paciente y elimina las etiquetas privadas, preservando los datos de imagen. La tabla "
        "siguiente delimita las operaciones del código, con independencia del contenido concreto "
        "de las muestras técnicas utilizadas para comprobarlo."
    )
    add_table(
        doc,
        ["Sí realiza", "No realiza"],
        [
            ("Sustitución de PatientName.", "Sustitución o eliminación de PatientID."),
            ("Eliminación global de etiquetas privadas.", "Supresión de fecha de nacimiento, sexo, fechas del estudio o institución."),
            ("Conservación de PixelData.", "Remapeo de StudyInstanceUID, SeriesInstanceUID o SOPInstanceUID."),
            ("Escritura de una copia DICOM.", "Procesamiento recursivo de carpetas o generación de registro de auditoría."),
            ("", "Detección/eliminación de texto que pudiera estar incrustado en los píxeles."),
        ],
    )
    doc.add_heading("6. Relación con el proceso de curado del entregable E6", level=1)
    doc.add_paragraph(
        "La pseudoanonimización de los datos se realizó mediante la ejecución del script de "
        "pseudoanonimización en el flujo de trabajo de la plataforma DIPCAN, por lo que Artelnics "
        "nunca tuvo acceso a nombres, documentos de identidad, datos de contacto ni otros "
        "identificadores directos de los pacientes."
    )
    doc.add_paragraph(
        "De esta manera, todos los conjuntos de datos utilizados por Artelnics para la construcción, "
        "el entrenamiento y la validación de los modelos se encontraban previamente pseudoanonimizados "
        "e identificados con el código de estudio DIPCANXXXXX."
    )
    doc.add_paragraph(
        "El script constituye una operación previa a la extracción de características radiómicas: "
        "recibe la imagen DICOM, elimina metadatos privados, sustituye el nombre y conserva el contenido "
        "de imagen. Las características cuantitativas obtenidas posteriormente pueden vincularse al "
        "registro de estudio usado por la plataforma, sin necesidad de incorporar el nombre del paciente "
        "a la matriz de datos curada."
    )
    doc.add_paragraph(
        "A efectos del requerimiento, se aportan conjuntamente el código fuente original, el par de "
        "archivos de ejemplo antes/después y la presente verificación. La custodia de las correspondencias "
        "entre códigos de estudio e identidad clínica, si existiera, corresponde al entorno que ejecutó "
        "el proceso y no forma parte del script entregado."
    )

    doc.add_heading("7. Inventario de evidencias", level=1)
    add_table(
        doc,
        ["Evidencia", "Descripción"],
        [
            ("dicom_anonymizer/dicom_anonymizer.py", "Código fuente original, 20 líneas."),
            ("dicom_anonymizer/example.dcm", "Archivo DICOM de entrada incluido con el script."),
            ("dicom_anonymizer/anonymized_example.DCM", "Resultado producido por el script."),
            ("Presente documento", "Descripción fiel, modo de uso, resultados y límites observados."),
        ],
    )
    doc.add_paragraph(
        "Nota de protección de datos: la verificación se limita al par de archivos de ejemplo "
        "incluido con el script. Los valores de cabecera no se reproducen: se expresan únicamente "
        "como estados, longitudes y resultados de comparación."
    )
    doc.add_paragraph(
        "[1] pydicom: biblioteca de código abierto para la lectura, modificación y escritura de "
        "archivos DICOM desde Python (https://pydicom.github.io/)."
    )

    doc.core_properties.title = "DIPCAN E6 — Script de anonimización DICOM"
    doc.core_properties.subject = "Respuesta al requerimiento técnico TSI-100206-2021-2"
    doc.core_properties.author = "Artificial Intelligence Techniques, S.L. (Artelnics)"
    doc.core_properties.comments = "Actualizado para incorporar el script original y ejemplos verificados."
    doc.save(OUTPUT)
    print(OUTPUT)
    print(BACKUP)


if __name__ == "__main__":
    main()
