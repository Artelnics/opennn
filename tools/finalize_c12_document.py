from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = Path(r"C:\Users\Roberto\OneDrive - artelnics.com\DIPCAN\subsanacion_tecnica")
DOCX = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO.docx"
BACKUP = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO_pre_finalizacion.bak.docx"

def find_one(doc, prefix):
    matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}; found {len(matches)}")
    return matches[0]

def insert_after(paragraph, text, style="Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    result.style = style
    result.add_run(text)
    return result

def main():
    copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    p = find_one(doc, "Quibim realizó una extracción bruta de 1.018")
    p.text = (
        "Quibim realizó una extracción bruta de 1.018 características radiómicas por lesión. "
        "Posteriormente, durante el curado y la agregación de la información por paciente, este "
        "conjunto se transformó en 127 variables radiómicas curadas, cifra verificada en la base "
        "de datos curada del proyecto. Sobre estas variables se aplicó el proceso de selección "
        "para reducir redundancias y correlaciones antes del entrenamiento del modelo."
    )

    p = find_one(doc, "La extracción inicial realizada por Quibim produjo 1.018")
    p.text = (
        "La extracción inicial realizada por Quibim produjo 1.018 características radiómicas "
        "brutas por lesión. Tras el curado y la agregación de la información por paciente, la "
        "base de datos curada del proyecto contiene 127 variables radiómicas. El análisis de "
        "redundancia y correlación redujo posteriormente este conjunto a 74 variables empleadas "
        "como entradas del modelo, tal como se recoge en la sección de métodos."
    )

    count_table = next(
        table
        for table in doc.tables
        if any("1.018" in cell.text for row in table.rows for cell in row.cells)
    )
    curated_row = next(
        row for row in count_table.rows if "Variables radiómicas curadas" in row.cells[0].text
    )
    curated_row.cells[0].text = "Variables radiómicas curadas"

    heading = find_one(doc, "Variables eliminadas de radiómica")
    heading.text = "Selección de variables eliminadas de radiómica"
    insert_after(
        heading,
        "El siguiente listado recoge una selección documentada de variables eliminadas durante "
        "las distintas etapas de depuración y selección. No constituye una relación exhaustiva "
        "de las 53 variables descartadas entre el conjunto curado de 127 variables y las 74 "
        "entradas finales del modelo.",
    )

    doc.core_properties.title = "Anexo C12. Desarrollo y validación clínica de algoritmos"
    doc.core_properties.subject = "Proyecto DIPCAN — versión final"
    doc.core_properties.comments = (
        "Versión final: 1.018 características brutas por lesión, 127 variables curadas, "
        "74 entradas del modelo y listado de eliminadas identificado como selección no exhaustiva."
    )
    doc.save(DOCX)
    print(DOCX)
    print(BACKUP)

if __name__ == "__main__":
    main()
