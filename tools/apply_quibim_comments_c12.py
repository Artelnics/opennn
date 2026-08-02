from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


BASE = Path(r"C:\Users\Roberto\OneDrive - artelnics.com\DIPCAN\subsanacion_tecnica")
DOCX = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO.docx"
BACKUP = BASE / "Anexo C12. Desarrollo y validación clínica de algoritmos_NUEVO_pre_comentarios_Quibim.bak.docx"


REPLACEMENTS = {
    "Por ello, planteamos un modelo que predice si existe una mutación en un gen a partir de valores extraídos de imágenes radiómicas.": (
        "Por ello, planteamos un modelo que predice si existe una mutación en un gen a partir de las "
        "características radiómicas extraídas de las imágenes de RM. Esto permitiría identificar "
        "mutaciones en un tumor mediante métodos no invasivos."
    ),
    "El objetivo es estimar la presencia o ausencia de una mutación genética mediante un método menos invasivo y más común, como la imagen radiómica.": (
        "El objetivo es estimar la presencia o ausencia de una mutación genética mediante el estudio "
        "de las características radiómicas extraídas de las imágenes de RM, utilizando un método "
        "menos invasivo y basado en exploraciones de uso habitual."
    ),
    "Las imágenes radiómicas muestran los órganos y tejidos del cuerpo": (
        "Las características radiómicas extraídas de las imágenes de RM son variables cuantitativas "
        "obtenidas de una imagen médica mediante algoritmos matemáticos. Describen propiedades de una "
        "región de interés (ROI), como un tumor o un órgano, más allá de lo que puede apreciarse visualmente."
    ),
    "La siguiente figura muestra varias imágenes de radiómica.": (
        "La siguiente figura muestra varias imágenes de RM con las segmentaciones de las lesiones."
    ),
    "Ejemplos de imágenes radiómicas.": (
        "Ejemplos de imágenes de RM con las segmentaciones de las lesiones medibles identificadas."
    ),
    "En este caso, modelizamos la probabilidad de que un paciente sea positivo para mutaciones": (
        "En este caso, modelizamos la probabilidad de que un paciente sea positivo para mutaciones "
        "en TP53, KRAS, PIK3CA, ERBB2, EGFR o PTEN, condicionada a las características radiómicas "
        "extraídas de las imágenes de RM."
    ),
    "A partir de las imágenes radiómicas, el equipo de Quibim genera": (
        "A partir de las secuencias de difusión de resonancia magnética adquiridas por el Hospital "
        "MD Anderson, el equipo de Quibim llevó a cabo la segmentación de las lesiones medibles "
        "previamente identificadas por el servicio de Radiología en la documentación complementaria. "
        "Sobre las regiones tumorales segmentadas se realizó posteriormente la extracción de "
        "características radiómicas, permitiendo obtener biomarcadores cuantitativos derivados de imagen."
    ),
    "Por otro lado, la predicción de la mutación de un gen a partir de imágenes radiómicas": (
        "Por otro lado, la predicción de la mutación de un gen a partir de características radiómicas "
        "extraídas de imágenes de RM puede resultar beneficiosa para los oncólogos. Esto les permite "
        "seleccionar medicamentos específicos y personalizados para cada paciente, lo que contribuye "
        "a mejorar su salud y bienestar."
    ),
    "Al combinar información de varios grupos, se busca obtener una descripción más completa": (
        "Al combinar información de varios grupos, se busca obtener una descripción más completa de "
        "la textura, la forma y otras propiedades relevantes derivadas de las imágenes de RM, con fines "
        "de diagnóstico y caracterización de tejidos."
    ),
    "En primer lugar, contábamos con 128 variables que definían una imagen radiómica.": (
        "Tras el curado y la agregación de la información por paciente, se dispuso de 127 variables "
        "radiómicas candidatas. El análisis de redundancia y correlación redujo este conjunto a 74 "
        "variables empleadas como entradas del modelo, tal como se recoge en la sección de métodos."
    ),
}


ADC_TEXT = (
    "Cuando las secuencias DWI disponían de un valor b bajo y otro alto, así como de los metadatos "
    "necesarios, también se calcularon cinco características relacionadas con el coeficiente de difusión "
    "aparente (ADC): media, mediana, desviación estándar y percentiles 25 y 75. Debido a que estos "
    "requisitos no se cumplían en todas las adquisiciones, las características ADC no pudieron obtenerse "
    "para todos los pacientes y no se utilizaron en el modelo final."
)

PYRADIOMICS_TEXT = (
    "Para elaborar este anexo, Artelnics se ha basado en el diccionario de variables radiómicas "
    "extraídas por Quibim y en la documentación de la biblioteca PyRadiomics "
    "(https://pyradiomics.readthedocs.io/en/latest/features.html). Se mantiene a continuación el "
    "listado detallado de características para facilitar la trazabilidad de las variables empleadas."
)

CLINICAL_BENEFIT_TEXT = (
    "Tras una validación clínica adicional, este enfoque podría contribuir a reducir la necesidad de "
    "determinados procedimientos invasivos, como la biopsia, al aprovechar imágenes de RM ya adquiridas "
    "durante la atención del paciente."
)


def replace_paragraph(paragraph, new_text):
    # The affected paragraphs use a uniform Normal style. Replacing their text preserves the
    # paragraph-level style, numbering and location while avoiding partial-run artifacts.
    paragraph.text = new_text


def insert_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def main():
    copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    applied = []
    for prefix, new_text in REPLACEMENTS.items():
        matches = [p for p in doc.paragraphs if p.text.startswith(prefix)]
        if len(matches) != 1:
            raise RuntimeError(f"Expected exactly one paragraph for {prefix!r}; found {len(matches)}")
        replace_paragraph(matches[0], new_text)
        applied.append(prefix)

    method_para = next(
        p for p in doc.paragraphs
        if p.text.startswith("A partir de las secuencias de difusión de resonancia magnética")
    )
    insert_after(method_para, ADC_TEXT, "Normal")

    conclusion_para = next(
        p for p in doc.paragraphs
        if p.text.startswith("Por otro lado, la predicción de la mutación de un gen")
    )
    insert_after(conclusion_para, CLINICAL_BENEFIT_TEXT, "Normal")

    annex_heading = next(
        p for p in doc.paragraphs
        if p.text.startswith("Anexo 3. Definición de características radiómicas")
    )
    insert_after(annex_heading, PYRADIOMICS_TEXT, "Normal")

    doc.core_properties.comments = (
        "Comentarios pertinentes de Quibim aplicados: terminología RM/radiómica, flujo de "
        "segmentación y extracción, difusión/ADC, referencia PyRadiomics y beneficio clínico prudente. "
        "Listado del Anexo 3 conservado."
    )
    doc.save(DOCX)
    print(DOCX)
    print(BACKUP)
    print(f"Replacements: {len(applied)}; inserted paragraphs: 3")


if __name__ == "__main__":
    main()
