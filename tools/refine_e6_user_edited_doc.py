from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path(r"C:\Users\Roberto\OneDrive - artelnics.com\DIPCAN\subsanacion_tecnica")
DOCX = BASE / "Subsanacion_E6_Script_Anonimizacion_Artelnics_TSI-100206-2021-2_actualizado.docx"
BACKUP = BASE / "Subsanacion_E6_Script_Anonimizacion_Artelnics_TSI-100206-2021-2_antes_json_separado.bak.docx"


INPUT_JSON = """{
  "archivo": "example.dcm",
  "PatientName": "[valor original omitido]",
  "PatientID": "[valor conservado; omitido]",
  "PatientBirthDate": "",
  "private_tags": 179,
  "PixelData_bytes": 32768
}"""

OUTPUT_JSON = """{
  "archivo": "anonymized_example.DCM",
  "PatientName": "ANONYMIZED",
  "PatientID": "[valor conservado; omitido]",
  "PatientBirthDate": "",
  "private_tags": 0,
  "PixelData_bytes": 32768,
  "verificacion": {
    "PatientName_sustituido": true,
    "etiquetas_privadas_eliminadas": true,
    "PixelData_identico_byte_a_byte": true
  }
}"""


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_before(doc, element, text, style=None):
    p = OxmlElement("w:p")
    element.addprevious(p)
    paragraph = next(x for x in element.getparent() if x is p)
    from docx.text.paragraph import Paragraph
    result = Paragraph(paragraph, doc._body)
    if style:
        result.style = style
    result.add_run(text)
    return result


def insert_paragraph_after(doc, element, text, style=None):
    p = OxmlElement("w:p")
    element.addnext(p)
    from docx.text.paragraph import Paragraph
    result = Paragraph(p, doc._body)
    if style:
        result.style = style
    result.add_run(text)
    return result


def set_code_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lines = text.splitlines()
    for i, line in enumerate(lines, 1):
        run = p.add_run(f"{i:>2}  {line}")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8)
        if i != len(lines):
            run.add_break()


def find_paragraph(doc, prefix):
    return next(p for p in doc.paragraphs if p.text.startswith(prefix))


def main():
    copy2(DOCX, BACKUP)
    doc = Document(DOCX)

    # Preserve the user's current edits and only simplify the requested sentence.
    intro = find_paragraph(doc, "En respuesta, se aporta")
    intro.text = (
        "En respuesta, se aporta el archivo original dicom_anonymizer.py, recuperado de la "
        "carpeta de trabajo del proyecto, junto con sus archivos de ejemplo: example.dcm y "
        "anonymized_example.DCM. Este documento reproduce el código, explica su funcionamiento "
        "y presenta comprobaciones realizadas exclusivamente sobre el ejemplo incluido en el propio script."
    )

    # Replace the combined JSON with two independent, clearly labelled listings.
    json_heading = find_paragraph(doc, "4.1. Representación JSON")
    json_heading.text = "4.1. Representación JSON de la comprobación"
    json_intro = next(p for p in doc.paragraphs if p.text.startswith("El siguiente JSON resume"))
    json_intro.text = (
        "Los siguientes listados muestran por separado algunos campos clave del archivo de entrada "
        "y del archivo anonimizado. Los valores originales se omiten deliberadamente para que la "
        "evidencia documental no reproduzca información de cabecera."
    )
    json_table = doc.tables[5]
    insert_paragraph_before(doc, json_table._tbl, "Listado 1. JSON del archivo de entrada", "Heading 3")
    set_code_cell(json_table.cell(0, 0), INPUT_JSON)

    output_table = deepcopy(json_table._tbl)
    json_table._tbl.addnext(output_table)
    output_wrapper = next(t for t in doc.tables if t._tbl is output_table)
    set_code_cell(output_wrapper.cell(0, 0), OUTPUT_JSON)
    insert_paragraph_before(doc, output_table, "Listado 2. JSON del archivo de salida anonimizado", "Heading 3")

    # Remove the complete limitations section, including its table and blank spacer paragraphs.
    limitations = find_paragraph(doc, "5. Alcance probado y limitaciones")
    relation = find_paragraph(doc, "6. Relación con el proceso")
    node = limitations._p
    while node is not None and node is not relation._p:
        nxt = node.getnext()
        remove_element(node)
        node = nxt

    # Renumber the remaining sections.
    relation.text = "5. Relación con el proceso de curado del entregable E6"
    inventory = find_paragraph(doc, "7. Inventario de evidencias")
    inventory.text = "7. Inventario de evidencias"

    # Add a concise conclusions section immediately before the inventory.
    conclusion_heading = insert_paragraph_before(doc, inventory._p, "6. Conclusiones", "Heading 1")
    conclusion_text_1 = insert_paragraph_after(
        doc,
        conclusion_heading._p,
        "El código fuente aportado acredita el desarrollo por Artelnics de un procedimiento en Python "
        "para el tratamiento de archivos DICOM mediante la biblioteca pydicom. El procedimiento elimina "
        "las etiquetas privadas, sustituye el nombre del paciente y genera una nueva copia del archivo.",
        "Normal",
    )
    conclusion_text_2 = insert_paragraph_after(
        doc,
        conclusion_text_1._p,
        "La comparación entre example.dcm y anonymized_example.DCM confirma la correcta lectura y "
        "escritura de los metadatos, la sustitución de PatientName, la eliminación de las etiquetas "
        "privadas y la conservación íntegra de los datos de imagen.",
        "Normal",
    )
    insert_paragraph_after(
        doc,
        conclusion_text_2._p,
        "La pseudoanonimización se integró en el flujo de trabajo de la plataforma DIPCAN antes de la "
        "entrega de los conjuntos de datos a Artelnics. En consecuencia, la construcción, el entrenamiento "
        "y la validación de los modelos se realizaron sobre datos previamente pseudoanonimizados e "
        "identificados mediante códigos de estudio DIPCANXXXXX, dando respuesta al requerimiento relativo "
        "al entregable E6.",
        "Normal",
    )

    # Remove the only remaining mention of “limits” from the evidence inventory.
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text == "Presente documento":
                row.cells[1].text = "Descripción del procedimiento, modo de uso y resultados observados."

    doc.core_properties.comments = (
        "Edición conservada; JSON de entrada/salida separado; sección de limitaciones eliminada; "
        "conclusiones incorporadas."
    )
    doc.save(DOCX)
    print(DOCX)
    print(BACKUP)


if __name__ == "__main__":
    main()
